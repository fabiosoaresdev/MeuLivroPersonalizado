from pyautogui import hotkey
from time import sleep
import os, subprocess, shutil
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx2pdf import convert

continuar='S'
while continuar != 'N':
    def l():
        print('\033[32m-=-\033[m' * 20)


    def substituir_palavras(paragrafo, palavras_substituir):
        for busca, substituicao in palavras_substituir.items():
            if busca in paragrafo.text:
                for run in paragrafo.runs:
                    run.text = run.text.replace(busca, substituicao)

    def obter_sexo():
        while True:
            sexo = input('INFORME O SEXO DA CRIANÇA [M/F]: ').upper().strip()
            if sexo in ['M', 'F']:
                return sexo
            else:
                print('Opção inválida. Por favor, informe M para masculino ou F para feminino.')
    l()
    print(' '*25 ,'VAMOS LÁ!')
    l()
    nome = input('INFORME O NOME DA CRIANÇA: ').title()
    sobrenome = input('AGORA, INFORME O SOBRENOME DA CRIANÇA: ').title()
    qmofrc = input('INFORME QUEM OFERECE O LIVRO: ').title()
    amg1 = input('INFORME O NOME DO PRIMEIRO AMIGO(A): ').title()
    amg2 = input('INFORME O NOME DO SEGUNDO AMIGO(A): ').title()
    amg3 = input('INFORME O NOME DO TERCEIRO AMIGO(A): ').title()
    cidade = input('INFORME A CIDADE DA CRIANÇA: ').title()
    dedicatoria = input('DIGITE A DEDICATÓRIA DO LIVRO: ').title()
    
    # Loop para confirmar se os dados estão corretos
    confirmado = False
    while not confirmado:
        l()
        print(f'\nConfirme os dados: \nNome: {nome} \nSobrenome: {sobrenome} \nQuem oferece o livro: {qmofrc} \nAmigo 1: {amg1} \nAmigo 2: {amg2} \nAmigo 3: {amg3} \nCidade: {cidade} \nDedicatória: {dedicatoria}')
        confirmar = input('Os dados estão corretos? [S/N] ').upper()
        if confirmar == 'S':
            confirmado = True
        elif confirmar == 'N':
            l()
            erro = int(input('Qual dado está errado? \n1. Nome \n2. Sobrenome \n3. Quem oferece o livro \n4. Amigo 1 \n5. Amigo 2 \n6. Amigo 3 \n7. Cidade \n8. Dedicatória \nDigite o número correspondente: '))
            if erro == 1:
                nome = input('Digite o nome correto: ').title()
            elif erro == 2:
                sobrenome = input('Digite o sobrenome correto: ').title()
            elif erro == 3:
                qmofrc = input('Digite quem oferece o livro correto: ').title()
            elif erro == 4:
                amg1 = input('Digite o nome do primeiro amigo correto: ').title()
            elif erro == 5:
                amg2 = input('Digite o nome do segundo amigo correto: ').title()
            elif erro == 6:
                amg3 = input('Digite o nome do terceiro amigo correto: ').title()
            elif erro == 7:
                cidade = input('Digite a cidade correta: ').title()
            elif erro == 8:
                dedicatoria = input('Digite a dedicatória correta: ').title()
            else:
                print('Opção inválida. Tente novamente.')

    l()
    sexo = obter_sexo()

    if sexo == 'M':
        pasta = 'Livros Meninos'
    elif sexo == 'F':
        pasta = 'Livros Meninas'
   
    # Lista todos os arquivos na pasta selecionada
    files = os.listdir(pasta)

    # Filtra apenas os arquivos com extensão .docx
    docx_files = list(filter(lambda x: x.endswith('.docx'), files))

    if len(docx_files) == 0:
        print('Nenhum arquivo .docx encontrado na pasta selecionada.')
    else:
        print('ESCOLHA O LIVRO DESEJADO:')
        l()
        for i, file in enumerate(docx_files):
            print(f'{i+1}. {file}')
        l()
        while True:
            escolha = input('   INFORME O NÚMERO CORRESPONDENTE AO LIVRO DESEJADO: ')
            if escolha.isdigit() and int(escolha) <= len(docx_files):
                break
            else:
                print('Opção inválida. Por favor, escolha um número válido.')

    # Obtém o arquivo escolhido com base na escolha do usuário
    chosen_file = docx_files[int(escolha) - 1]
    print(f'Arquivo escolhido: {chosen_file}')

    # Abre o arquivo .docx selecionado
    doc = Document(os.path.join(pasta, chosen_file))

    # Copia o arquivo .docx para um novo arquivo chamado "livro cópia.docx"
    doc_copy_path = os.path.join(pasta, 'livro cópia.docx')
    shutil.copy2(os.path.join(pasta, chosen_file), doc_copy_path)

    # Renomeia o arquivo com o novo nome
    os.rename(doc_copy_path, os.path.join(pasta, 'Livro Cópia.docx'))

    # Abre a cópia do arquivo .docx
    doc_copy = Document(doc_copy_path)

    # Função para substituir palavras no documento
    def substituir_palavra(paragrafo, palavra_antiga, palavra_nova):
        if palavra_antiga in paragrafo.text:
            inline = paragrafo.runs
            for i in range(len(inline)):
                if palavra_antiga in inline[i].text:
                    texto = inline[i].text.replace(palavra_antiga, palavra_nova)
                    inline[i].text = texto

    # Substitui as palavras no documento
    for paragrafo in doc_copy.paragraphs:
        
        substituir_palavra(paragrafo, 'PERSONAGEM', nome)
        substituir_palavra(paragrafo, 'Personagem', nome)
        substituir_palavra(paragrafo, '<<nome>>', nome)
        substituir_palavra(paragrafo, 'SOBRENOME', sobrenome)
        substituir_palavra(paragrafo, 'DEDICATORIA', dedicatoria)
        substituir_palavra(paragrafo, 'QUEM OFERECE', qmofrc)
        substituir_palavra(paragrafo, 'CIDADE', cidade)
        substituir_palavra(paragrafo, 'AMIGO 01', amg1)
        substituir_palavra(paragrafo, 'AMIGO 02', amg2)
        substituir_palavra(paragrafo, 'AMIGO 03', amg3)
        substituir_palavra(paragrafo, 'sexo', sexo)
        
        # Define a cor da fonte para preto
    for paragrafo in doc_copy.paragraphs:
        for run in paragrafo.runs:
            font_color = run.font.color
            if font_color is not None and font_color.rgb != RGBColor(0, 0, 0):
                run.font.color.rgb = RGBColor(0, 0, 0)

    # Salva o arquivo .docx com as palavras substituídas
    doc_copy_name = chosen_file.split(f'História de{nome} {sobrenome}')[0] + '.docx'
    doc_copy_path = os.path.join(pasta, doc_copy_name)
    doc_copy.save(doc_copy_path)

    # Renomeia o arquivo gerado
    novo_nome = 'novo_nome.docx'
    copia_path = os.path.join(pasta, novo_nome)
    os.rename(doc_copy_path, copia_path)

    # Converte o arquivo .docx para .pdf
    pdf_path = os.path.join('livros prontos', f'{nome}_{sobrenome}.pdf')
    convert(copia_path, pdf_path)

    # Abre o arquivo .pdf
    subprocess.run(['start', '', pdf_path], shell=True)

    sleep(1.2)
    hotkey('ctrl','p')
    
    # Remove o arquivo .docx gerado
    os.remove(copia_path)
        
    # Remove o arquivo .docx de ambos os diretórios (Livro Meninas e Livro Meninos)
    caminho_arquivo_meninas = os.path.join(os.getcwd(), 'Livro Meninas', novo_nome)
    caminho_arquivo_meninos = os.path.join(os.getcwd(), 'Livro Meninos', novo_nome)

    if os.path.exists(caminho_arquivo_meninas):
        os.remove(caminho_arquivo_meninas)
        print(f'O arquivo {novo_nome} foi excluído do diretório "Livro Meninas".')

    if os.path.exists(caminho_arquivo_meninos):
        os.remove(caminho_arquivo_meninos)
        print(f'O arquivo {novo_nome} foi excluído do diretório "Livro Meninos".')


    l()
    print(f'Arquivo PDF salvo em: {pdf_path}')

    continuar=input('Você deseja fazer um novo livro? [S/N]').upper()





